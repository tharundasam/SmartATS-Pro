"""
Raw text extraction from resume files.

PDF: pdfplumber first (better layout/whitespace handling for resumes,
which are often multi-column or table-like), falling back to pypdf if
pdfplumber raises on a malformed/unusual PDF. DOCX: python-docx, which
reads paragraphs and table cells (resumes sometimes use tables for
layout, e.g. a skills grid).

This module only extracts raw text — it has no opinion about names,
emails, skills, etc. That logic lives in rule_based_extractor.py, kept
separate so the "get text out of a file" concern and the "find
structured fields in that text" concern can be tested and changed
independently.
"""

import io

import pdfplumber
import pypdf
from docx import Document
from fastapi import HTTPException, status


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Tries pdfplumber first; if it raises (corrupt/unusual PDF structure),
    falls back to pypdf rather than failing the whole request outright.
    """
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages_text).strip()
            if text:
                return text
            # pdfplumber opened the file but found no extractable text
            # (e.g. a scanned/image-only PDF) — fall through to pypdf,
            # which occasionally succeeds where pdfplumber returns empty.
    except Exception:
        pass

    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages_text = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages_text).strip()
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not extract text from this PDF. It may be a "
            "scanned image without selectable text, password-protected, "
            "or corrupted.",
        ) from exc

    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="No extractable text found in this PDF. It may be a "
            "scanned image without selectable text (OCR is not supported "
            "yet).",
        )
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts paragraph text and table-cell text (resumes sometimes use
    tables for layout, e.g. a two-column skills section)."""
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not read this DOCX file. It may be corrupted "
            "or not a valid Word document.",
        ) from exc

    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    text = "\n".join(parts).strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="No extractable text found in this DOCX file.",
        )
    return text


def extract_text(file_bytes: bytes, extension: str) -> str:
    """Dispatches to the right extractor based on file extension.
    `extension` is expected to already be validated (see
    app/utils/file_validation.py) — this function trusts it."""
    if extension == "pdf":
        return extract_text_from_pdf(file_bytes)
    if extension == "docx":
        return extract_text_from_docx(file_bytes)
    raise ValueError(f"Unsupported extension for text extraction: {extension}")
